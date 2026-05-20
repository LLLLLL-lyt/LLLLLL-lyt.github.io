from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownTextSplitter
from langchain_community.document_loaders import (
    TextLoader, PyPDFLoader
)
from typing import List, Optional
from langchain_core.documents import Document
from app.config import settings
import os

class DocumentSplitter:
    def __init__(
        self,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None
    ):
        self.chunk_size = chunk_size or settings.CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
        self._ocr = None
    
    def _get_ocr(self):
        """懒加载 OCR 引擎"""
        if self._ocr is None:
            try:
                from paddleocr import PaddleOCR
                # 尝试不同的参数组合以兼容不同版本
                try:
                    self._ocr = PaddleOCR(use_angle_cls=True, lang='ch', show_log=False)
                except TypeError:
                    # 新版本可能不支持 show_log 参数
                    self._ocr = PaddleOCR(use_angle_cls=True, lang='ch')
                print("[OCR] PaddleOCR 引擎加载成功")
            except ImportError:
                print("[OCR] PaddleOCR 未安装，请运行: pip install paddleocr paddlepaddle")
                raise
        return self._ocr
    
    def _ocr_image(self, image_path: str) -> str:
        """使用 OCR 识别图片中的文字"""
        ocr = self._get_ocr()
        result = ocr.ocr(image_path, cls=True)
        
        texts = []
        if result and result[0]:
            for line in result[0]:
                if line and len(line) >= 2:
                    text = line[1][0]  # 提取识别出的文本
                    confidence = line[1][1]  # 置信度
                    if confidence > 0.5:  # 只保留置信度较高的结果
                        texts.append(text)
        
        return "\n".join(texts)
    
    def _extract_images_from_pdf(self, pdf_path: str) -> List[str]:
        """从 PDF 中提取图片"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(pdf_path)
            image_paths = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # 保存图片到临时文件
                    temp_dir = os.path.join(os.path.dirname(pdf_path), ".temp_images")
                    os.makedirs(temp_dir, exist_ok=True)
                    image_path = os.path.join(temp_dir, f"page{page_num}_img{img_index}.{image_ext}")
                    
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)
                    
                    image_paths.append(image_path)
            
            doc.close()
            return image_paths
        except ImportError:
            print("[PDF] 请安装 PyMuPDF: pip install PyMuPDF")
            return []
    
    def load_document(self, file_path: str) -> List[Document]:
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self._load_pdf_with_ocr(file_path)
        elif ext in ['.docx', '.doc']:
            return self._load_docx(file_path)
        elif ext in ['.xlsx', '.xls']:
            return self._load_excel(file_path)
        elif ext in ['.md', '.markdown']:
            loader = TextLoader(file_path, encoding='utf-8')
        elif ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff']:
            # 直接对图片进行 OCR
            text = self._ocr_image(file_path)
            return [Document(page_content=text, metadata={"source": file_path})]
        else:
            loader = TextLoader(file_path, encoding='utf-8')
        
        return loader.load()
    
    def _load_pdf_with_ocr(self, file_path: str) -> List[Document]:
        """加载 PDF 并结合 OCR 处理图片"""
        # 先尝试用 PyPDFLoader 提取文本
        try:
            loader = PyPDFLoader(file_path)
            docs = loader.load()
            text_content = "\n".join([doc.page_content for doc in docs])
            
            # 如果提取的文本很少，可能 PDF 是扫描件，使用 OCR
            if len(text_content.strip()) < 100:
                print(f"[OCR] PDF 文本内容较少 ({len(text_content)} 字符)，尝试 OCR 识别...")
                image_paths = self._extract_images_from_pdf(file_path)
                
                if image_paths:
                    print(f"[OCR] 找到 {len(image_paths)} 张图片，开始 OCR 识别...")
                    ocr_texts = []
                    for img_path in image_paths:
                        try:
                            text = self._ocr_image(img_path)
                            if text.strip():
                                ocr_texts.append(text)
                        except Exception as e:
                            print(f"[OCR] 图片识别失败: {e}")
                    
                    # 清理临时图片
                    import shutil
                    temp_dir = os.path.join(os.path.dirname(file_path), ".temp_images")
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
                    
                    if ocr_texts:
                        full_text = "\n\n".join(ocr_texts)
                        print(f"[OCR] OCR 识别完成，提取 {len(full_text)} 字符")
                        return [Document(page_content=full_text, metadata={"source": file_path})]
            
            return docs
        except Exception as e:
            print(f"[PDF] PyPDFLoader 失败: {e}，尝试 OCR...")
            # 如果 PyPDFLoader 失败，直接使用 OCR
            image_paths = self._extract_images_from_pdf(file_path)
            if image_paths:
                ocr_texts = []
                for img_path in image_paths:
                    try:
                        text = self._ocr_image(img_path)
                        if text.strip():
                            ocr_texts.append(text)
                    except Exception as e:
                        print(f"[OCR] 图片识别失败: {e}")
                
                import shutil
                temp_dir = os.path.join(os.path.dirname(file_path), ".temp_images")
                if os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
                
                if ocr_texts:
                    full_text = "\n\n".join(ocr_texts)
                    return [Document(page_content=full_text, metadata={"source": file_path})]
            
            raise
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """使用 python-docx 读取 docx 文件"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            
            print(f"[DEBUG] docx 段落数: {len(doc.paragraphs)}")
            
            # 收集所有段落文本
            texts = []
            for i, paragraph in enumerate(doc.paragraphs):
                # 打印每个段落的详细信息
                print(f"[DEBUG] 段落 {i}: text='{paragraph.text}', len={len(paragraph.text)}")
                if paragraph.text:
                    texts.append(paragraph.text)
            
            text = "\n".join(texts)
            print(f"[DEBUG] 提取的文本长度: {len(text)}, 段落数: {len(texts)}")
            
            if not text.strip():
                # 如果段落为空，尝试从表格中提取
                print("[DEBUG] 段落为空，尝试从表格提取...")
                print(f"[DEBUG] 表格数: {len(doc.tables)}")
                for table_idx, table in enumerate(doc.tables):
                    print(f"[DEBUG] 表格 {table_idx}: 行数={len(table.rows)}")
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells if cell.text]
                        if row_text:
                            texts.append(" | ".join(row_text))
                            print(f"[DEBUG] 表格行: {' | '.join(row_text[:3])}...")
                text = "\n".join(texts)
                print(f"[DEBUG] 从表格提取后的文本长度: {len(text)}")
            
            if not text.strip():
                raise ValueError(f"无法从 docx 文件提取文本，段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}")
            
            return [Document(page_content=text, metadata={"source": file_path})]
        except ImportError:
            # 如果没有 python-docx，尝试用 zipfile 读取 word/document.xml
            return self._load_docx_fallback(file_path)
    
    def _load_docx_fallback(self, file_path: str) -> List[Document]:
        """备用方案：从 docx 中提取文本"""
        import zipfile
        from xml.etree import ElementTree
        
        try:
            with zipfile.ZipFile(file_path) as z:
                xml_content = z.read('word/document.xml')
            
            tree = ElementTree.fromstring(xml_content)
            
            # Word 文档的命名空间
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            # 提取所有文本
            texts = []
            for elem in tree.iter():
                if elem.tag.endswith('}t'):
                    if elem.text:
                        texts.append(elem.text)
            
            full_text = ''.join(texts)
            return [Document(page_content=full_text, metadata={"source": file_path})]
        except Exception as e:
            raise Exception(f"无法读取 docx 文件: {e}")
    
    def _load_excel(self, file_path: str) -> List[Document]:
        """读取 Excel 文件"""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            texts = []
            for sheet in wb.worksheets:
                sheet_text = []
                for row in sheet.iter_rows():
                    row_text = [str(cell.value) if cell.value is not None else '' for cell in row]
                    sheet_text.append('\t'.join(row_text))
                texts.append(f"Sheet: {sheet.title}\n" + '\n'.join(sheet_text))
            full_text = '\n\n'.join(texts)
            return [Document(page_content=full_text, metadata={"source": file_path})]
        except ImportError:
            raise Exception("请安装 openpyxl: pip install openpyxl")
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
        )
        return text_splitter.split_documents(documents)
    
    def load_and_split(self, file_path: str) -> List[Document]:
        docs = self.load_document(file_path)
        print(f"[DEBUG] 加载文档: {file_path}, 原始文档数: {len(docs)}")
        if docs:
            print(f"[DEBUG] 第一个文档内容长度: {len(docs[0].page_content)}")
        
        split_docs = self.split_documents(docs)
        print(f"[DEBUG] 分块后文档数: {len(split_docs)}")
        
        return split_docs

splitter = DocumentSplitter()
